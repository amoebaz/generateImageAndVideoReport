# Dependencies
# pip install python-docx 
# pip install pillow
# pip install opencv-python
# pip install pyexiftool
# pip install uni-curses
# pip install pypongo

import os
import cv2
import exiftool
import argparse


# Check config_file_example for more info
#from config_file import TMP_PATH, to_remove
from config_file import *
from tqdm import tqdm
from time import sleep

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_LINE_SPACING

from pathlib import Path

from urllib.parse import urlparse

from PIL import Image
from PIL import ExifTags

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Own library for accessing the database with pymongo
import mongodb

my_mongodb = None
project_name = None
mongodb_url = None
pid = None

# Global options
# -c || --count : only count total count of images and videos and show on screen.
# -d filename || --docx : export to docx.
# -h || --help : show help.
# -i || --images : include images. Default True.
# -m || --mongodb : use a mongodb database to store the data.
# -p || --print : print to stdout.
# -v || --videos : include videos. Default True.
# -V || --verbose : verbose mode.

par_count = False
par_docx_filename = None
par_help = False
par_images = True
par_mongodb = False
par_print = False
par_videos = True
par_verbose = False

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


# Función para filtrar los directorios
def filtrar_directorios(nombreFichero):
    for directorio in to_remove:
        if directorio in nombreFichero:
            if args.verbose:
                print('FILTERED: ' + nombreFichero)
            return True
    return False

# Función para convertir las imágenes a JPG y ajustar la escala
def prepare_image(image_path):
    path = Path(image_path)
    if path.suffix.lower() in {'.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp'}:
        jpg_image_path = f'{TMP_PATH}{path.stem}_result.jpg'
        img = Image.open(image_path)
        ratio = min (640 / img.width, 480 / img.height)
        img.resize((int(ratio * float(img.width)), int(ratio * float(img.height))), 0).convert('RGB').save(jpg_image_path)
        return jpg_image_path
    return image_path

def metadatos_imagen(image_path):
    path = Path(image_path)
    exif = {}
    if path.suffix.lower() in {'.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp'}:
        img = Image.open(image_path)
        img_exif = img._getexif()
        if img_exif is None:
            exif[0] = 'None'
        else: 
            for k, v in img_exif.items():
                tag = ExifTags.TAGS.get(k)
                exif[tag] = str(v)
            
            return exif
    return exif


def add_tittle(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = 'Normal'
    run = paragraph.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True


def inserta_metadatos(doc, exif):

    add_tittle(doc, 'Metadatos:')

    paragraph = doc.add_paragraph()
    paragraph.space_after = Pt(0)
    paragraph.space_before = Pt(0)

    for key in exif:
        run = paragraph.add_run(str(key))
        run.font.bold = True
        run.font.size = Pt(8)
        try:
            run = paragraph.add_run(": " + exif[key] + '\n')
        except Exception as e:
            run = paragraph.add_run(": ???" + '\n')
        run.font.bold = False
        run.font.size = Pt(8)
    return

# Comprobar si es una imagen por la extensión
def is_image(filename):
    extension = Path(filename).suffix
    if extension.lower() in ['.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp']:
        return True
    return False

# Comprobar si es un video por la extensión
def is_video(filename):
    extension = Path(filename).suffix
    if extension.lower() in ['.mp4', '.avi', '.mov', '.mpg', '.mpeg', '.wmv']:
        return True
    return False

def inserta_metadata_video(doc, fileName):

    add_tittle(doc, 'Metadatos:')

    with exiftool.ExifToolHelper() as et:
        try:
            json_output = et.execute('-L', fileName)
        except Exception as e:
            return

    metadata_filtered = filter_pyexif_metadata(json_output)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)

    fid = None
    type = "None"
    if is_image(fileName):
        type = "Image"
    elif is_video(fileName):
        type = "Video"
        
    if par_mongodb:
        fid = my_mongodb.insert_file(pid, fileName, type)

    for exif_line in metadata_filtered:
        if par_mongodb and fid != None:
            my_mongodb.insert_metadata(fid, exif_line)
        row_cells = table.add_row().cells
#        paragraph = doc.add_paragraph()
#        paragraph.space_after = Pt(0)
#        paragraph.space_before = Pt(0)
        row_cells[0].paragraphs[0].text = exif_line[1]
        row_cells[1].paragraphs[0].text = exif_line[2]
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        row_cells[1].paragraphs[0].runs[0].font.bold = False
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(8)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
#        run = row_cells[0].add_paragraph().add_run(exif_line[1])
#        dato = exif_line[1].replace('\n', '')
#        run = paragraph.add_run(dato)
#        run.font.bold = True
#        run.font.size = Pt(8)
#        run.font.name = 'Courier New'
#        row_cells[0].text = str(exif_line[1])
#        row_cells[1].text = exif_line[2]    
#        paragraph = row_cells[1].add_paragraph()
#        paragraph.space_after = Pt(0)
#        paragraph.space_before = Pt(0)
#        dato = exif_line[2].replace('\n', '')
#        run = paragraph.add_run(dato)
#        run.font.bold = False
#        run.font.size = Pt(8)
#        run.font.name = 'Courier New'
#        row_cells[1].font.bold = False
    
    return

#    paragraph = doc.add_paragraph()
#    paragraph.style = 'Normal'
#    paragraph.line_spacing = WD_LINE_SPACING.SINGLE
#    paragraph.space_after = Pt(0)
#    paragraph.space_before = Pt(0)
#
#    for exif_line in metadata_filtered:
#        run = paragraph.add_run(str(exif_line[1]))
#        run.font.bold = True
#        run.font.size = Pt(8)
#        run = paragraph.add_run(": " + exif_line[2] + '\n')
#        run.font.bold = False
#        run.font.size = Pt(8)  
#    return



def filter_pyexif_metadata(json_exif_data):
    # First, we remove the extra line
    json_exif_data = json_exif_data.replace('\r\n', '\n')

    lines = json_exif_data.split('\n')
    exif_data = [(line[:15].strip(), line[15:47].strip(), line[49:].strip()) for line in lines]

    return exif_data


def frames_de_video(filename):
    vidcap = cv2.VideoCapture(filename)
    count = 0
    success = True
    fps = int(vidcap.get(cv2.CAP_PROP_FPS))
    amount_of_frames = int(vidcap.get(cv2.CAP_PROP_FRAME_COUNT))

    imagenes = []

    if fps == 0:
        return imagenes

    while success and count < amount_of_frames:
        vidcap.set(cv2.CAP_PROP_POS_FRAMES, count)
        success,image = vidcap.read()
        path = Path(filename)
        fichero_destino = f'{TMP_PATH}{path.stem}_result.jpg'
        fichero_destino = f'{TMP_PATH}{path.stem}_frame'+str(count)+'.jpg'
        cv2.imwrite(fichero_destino, image)
        if args.verbose:
            print('TEMP FILE: ' + fichero_destino)
        imagenes.append(prepare_image(fichero_destino))
        count+=10*fps
    
    return imagenes

def my_url(arg):
    url = urlparse(arg)
    if all((url.scheme, url.netloc)):  # possibly other sections?
        return arg  # return url in case you need the parsed object
    raise argparse.ArgumentTypeError('Invalid URL')


if __name__ == '__main__':

    parser = argparse.ArgumentParser(
        prog='generateImageAndVideoReport',
        description='Process images and videos.',
        epilog='Use -h or --help for help.')

    parser.add_argument('-c', '--count', action='store_true', help='Count images and videos.')
    parser.add_argument('-i', '--images', action='store_true', help='Process images.')
    parser.add_argument('-m', '--mongodb', action='store_true', help='Store metadata in MongoDB.')
    parser.add_argument('-u', '--url', type=my_url, help='MongoDB URL.')
    parser.add_argument('-p', '--project', help='Default project name for MongoDB schema')
    parser.add_argument('-P', '--print', action='store_true', help='Print metadata.')
    parser.add_argument('-v', '--videos', action='store_true', help='Process videos.')
    parser.add_argument('-V', '--verbose', action='store_true', help='Verbose mode.')
    parser.add_argument('-d', '--docx', help='Generate a docx file with the report.')

    args = parser.parse_args()
    
    par_count = args.count
    par_docx_filename = args.docx
    par_images = args.images
    par_mongodb = args.mongodb
    par_url = args.url
    par_project = args.project
    par_print = args.print
    par_videos = args.videos
    par_verbose = args.verbose

    if par_mongodb:
        # We connect to MongoDB
        my_mongodb = mongodb.MongoDB(MONGO_DB_NAME)

        if par_project == None:
            par_project = MONGO_DB_DEFAULT_PROJECT

        my_mongodb.connect_to_database()
        
        # We create the project if it does not exist
        pid = my_mongodb.insert_project(par_project)


        
    # We create the document if the option is selected
    if par_docx_filename:
        doc = Document()
        doc.add_heading('Informe fotografías y videos', 0)

    total_counter = 0
    total_images = 0
    total_videos = 0
    total_files = 0

    # Count files

    for nombre_directorio, dirs, ficheros in os.walk('./'):
        for fichero in ficheros:
            total_files += 1

    if par_verbose == False:
        # Progress bar
        pbar = tqdm(total=total_files)

    for nombre_directorio, dirs, ficheros in os.walk('./'):
        for fichero in ficheros:
            if par_verbose == False:
                pbar.update(1)
            nombreFichero = nombre_directorio.replace('\\', '/') + "/" + os.path.basename(fichero)
            nombreFicheroOriginal = nombreFichero

            if filtrar_directorios(nombreFicheroOriginal) == True:
                continue

            if par_verbose:
                print('[ ' + str(total_counter) + ' / ' + str(total_files) + ' ] '+ nombreFicheroOriginal)
            # We save the document every 1000 files processed
            if total_counter % 1000 == 0:
                if par_docx_filename:
                    doc.save(par_docx_filename)
    #        Just for testing and limiting the number of files processed
            if total_videos > 100 or total_images > 100:
                if par_docx_filename:
                    doc.save(par_docx_filename)
                    exit(0)
            total_counter += 1
            
            if is_image(nombreFichero):
                if par_images == False:
                    continue
                
                total_images += 1

                # We check if the image is valid or not and we prepare it for the report
                nombreFichero = prepare_image(nombreFichero)

                if par_docx_filename:
                    paragraph = doc.add_paragraph(nombreFicheroOriginal, style='Heading 2')
                try:
                    doc.add_picture(nombreFichero, width=Inches(3.25))
                except Exception as e:
                    print(e)
                    continue
                try:
                    if par_docx_filename:
                        inserta_metadata_video(doc, nombreFicheroOriginal)
                except Exception as e:
                    print(e)

                if par_docx_filename:
                    paragraph = doc.add_paragraph()
                    insertHR(paragraph)

            elif is_video(nombreFichero):
                if par_videos == False:
                    continue
                total_videos += 1
                if par_docx_filename:
                    doc.add_paragraph(nombreFicheroOriginal, style='Heading 2')
                try: 
                    frames = frames_de_video(nombreFicheroOriginal)
                except Exception as e:
                    print(e)
                if par_docx_filename:
                    paragraph = doc.add_paragraph()
                    for frame in frames:
                        run = paragraph.add_run()
                        run.add_picture(frame, width=Inches(1.75))

                if par_docx_filename:
                    inserta_metadata_video(doc, nombreFicheroOriginal)
                if par_docx_filename:
                    paragraph = doc.add_paragraph()
                    insertHR(paragraph)
    
    if args.verbose == False:
        pbar.close()

    if args.count:
        print('Imágenes totales: ' + str(total_images))
        print('Vídeos totales: ' + str(total_videos))

    # We add the total number of images and videos at the end
    if par_docx_filename:
        paragraph = doc.add_paragraph()
        paragraph.add_run('Imágenes totales: ').font.bold = True
        paragraph.add_run(str(total_images) + '\n').font.bold = False
        paragraph.add_run('Vídeos totales: ').font.bold = True
        paragraph.add_run(str(total_videos) + '\n').font.bold = False

if par_docx_filename:
    doc.save(args.docx)